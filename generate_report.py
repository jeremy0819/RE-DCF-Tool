# -*- coding: utf-8 -*-
"""
generate_report.py — RE-DCF-Tool 專案報告生成器
執行：python generate_report.py
輸出：RE-DCF-Tool_專案報告_2026-06-26.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 品牌色彩系統 ────────────────────────────────────────────────────
BRAND_PURPLE   = RGBColor(0x53, 0x4A, 0xB7)   # #534AB7 主品牌色
DARK_NAVY      = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A 主文字
MID_GRAY       = RGBColor(0x47, 0x56, 0x69)   # #475569 次文字
LIGHT_GRAY     = RGBColor(0x94, 0xA3, 0xB8)   # #94A3B8 說明文字
TABLE_HEADER   = RGBColor(0x53, 0x4A, 0xB7)   # 表格標題背景同品牌色
TABLE_STRIPE   = RGBColor(0xF1, 0xF0, 0xFB)   # #F1F0FB 淺紫條紋
GREEN_OK       = RGBColor(0x16, 0xA3, 0x4A)   # ✅ 綠
AMBER_WARN     = RGBColor(0xD9, 0x77, 0x06)   # ⚠️ 黃橘
RED_RISK       = RGBColor(0xDC, 0x26, 0x26)   # 🔴 紅
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)


# ── 輔助函式 ─────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)          # RGBColor.__str__ 回傳 'RRGGBB' 六位 hex
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), val.get('val', 'single'))
            b.set(qn('w:sz'), str(val.get('sz', 4)))
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(b)
    tcPr.append(tcBorders)


def add_paragraph_border_bottom(para, color="534AB7", sz=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_section_heading(doc, text, level=1):
    """品牌化章節標題：品牌紫色底線"""
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_PURPLE
        add_paragraph_border_bottom(p, color="534AB7", sz=12)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = DARK_NAVY
        add_paragraph_border_bottom(p, color="94A3B8", sz=4)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = MID_GRAY
    return p


def add_body(doc, text, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK_NAVY
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = color or DARK_NAVY
    return p


def make_table(doc, headers, rows, col_widths=None, stripe=True):
    """標準報告表格：品牌紫表頭 + 交替條紋"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 欄寬
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # 表頭
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, TABLE_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    # 資料列
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if stripe and ri % 2 == 1:
            bg = TABLE_STRIPE
        else:
            bg = WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            # 支援 tuple (text, bold, color)
            if isinstance(val, tuple):
                text, bold, clr = val[0], val[1] if len(val) > 1 else False, val[2] if len(val) > 2 else DARK_NAVY
            else:
                text, bold, clr = str(val), False, DARK_NAVY
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(9)
            run.font.color.rgb = clr
    return table


def add_callout(doc, label, text, label_color=None):
    """側邊標注框：如風險、提示"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    lc = label_color or BRAND_PURPLE
    r1 = p.add_run(f"  {label}  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = WHITE
    # 用底線模擬 badge（Word 無法內嵌 HTML，改用粗體+顏色）
    r1.font.highlight_color = None
    r1.font.color.rgb = lc
    r2 = p.add_run(f"  {text}")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK_NAVY
    return p


# ══════════════════════════════════════════════════════════════════════
#  報告本體
# ══════════════════════════════════════════════════════════════════════

def build_report(research_notes: str = "") -> str:
    doc = Document()

    # ── 頁面設定 ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.top_margin   = Cm(2.0)
    section.bottom_margin= Cm(2.0)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 預設字體
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans TC')

    # ══════════════════════════════════════════════════════════════════
    # 封面區塊
    # ══════════════════════════════════════════════════════════════════
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(20)
    cover.paragraph_format.space_after  = Pt(4)
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cover.add_run("RE-DCF-Tool")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BRAND_PURPLE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub.paragraph_format.space_after = Pt(2)
    rs = sub.add_run("都更 / 危老前期評估工具　｜　專案技術報告")
    rs.font.size = Pt(13)
    rs.font.color.rgb = MID_GRAY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.paragraph_format.space_after = Pt(24)
    rm = meta.add_run("永盛開發建設　｜　報告日期：2026-06-26　｜　版本：v4.6")
    rm.font.size = Pt(9)
    rm.font.color.rgb = LIGHT_GRAY
    add_paragraph_border_bottom(meta, color="534AB7", sz=16)

    # ══════════════════════════════════════════════════════════════════
    # 1. 執行摘要
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "1  執行摘要", level=1)

    summary_rows = [
        ("定位", "內部 PropTech 工具，將 Excel 財務模型程式化，專為台灣都更/危老前期評估設計"),
        ("技術棧", "Python 3 · Streamlit ≥1.30 · Plotly ≥5.18 · Pandas ≥2.0（單一 repo，4 依賴）"),
        ("部署狀態", "✅  Streamlit Community Cloud 線上運行，main 分支推送即自動部署"),
        ("目前版本", "v4.6　｜　build 2026-06-23　｜　開發起始 2026-05-26"),
        ("程式規模", "4 模組共 1,825 行（app.py 1,170 / calc_engine.py 446 / law_db.py 124 / test_golden.py 85）"),
        ("計算驗證", "黃金測試 4 案全 PASS（安和段 / 龜山段 / 中正段 / 中正 L6 投報），容差 0.5 m²"),
        ("核心價值", "逐層 §162 查核自動偵測三項免計超出；共同負擔六科目完整拆解；L7 更新前估值"),
    ]
    make_table(doc, ["項目", "說明"], summary_rows, col_widths=[4, 12.5])

    doc.add_paragraph()
    add_body(doc,
        "本工具自 2026 年 5 月底起，以 27 次 commit 歷經 6 個主要版本迭代（v4.1–v4.6），"
        "完成從單一 app.py 到三層模組化（UI / 計算引擎 / 法規庫）的重構，"
        "並在第一個月內實現可部署、可測試、可供建築師與 PM 協作的可行性評估平台。",
        size=10)

    # ══════════════════════════════════════════════════════════════════
    # 2. 專案背景與目的
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "2  專案背景與目的", level=1)

    add_body(doc, "2.1  問題痛點", size=11, color=DARK_NAVY)
    bullets_bg = [
        "都更/危老前期評估高度依賴 Excel，多份試算表版本分散、公式易斷、難以協作稽核。",
        "逐層 §162 免計查核（梯廳/安全梯/陽台）若誤用「FA 總量法」，可導致陽台超出量放大約 10 倍（已知踩坑，v3 已修正）。",
        "共同負擔六大科目（都市更新權利變換實施辦法）無標準計算框架，不同人對費率基數（總銷 vs 工程費 vs 營造）認知不一致。",
        "建築師、代書、建設公司三方數字收斂前，PM 缺乏即時對帳工具。",
    ]
    for b in bullets_bg:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_body(doc, "2.2  目標使用者與場景", size=11, color=DARK_NAVY)
    user_rows = [
        ("🏗️ 建管確認", "容積合規狀態、三項免計查核結果、超出自動補計"),
        ("📐 建築師協作", "圖說面積表 Excel/CSV 匯入、逐層明細核對、§162 欄位對照卡"),
        ("📋 地政士（代書）", "容積帳數字、更新前土地現值估算（L7 §56 基準）"),
        ("💰 開發商 / PM", "共負比 / 報酬率 / 敏感度熱力圖 / 都更全案投報六科目"),
    ]
    make_table(doc, ["角色", "主要使用功能"], user_rows, col_widths=[4, 12.5])

    # ══════════════════════════════════════════════════════════════════
    # 3. 技術架構
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "3  技術架構", level=1)

    make_section_heading(doc, "3.1  六層計算骨架（L1–L7）", level=2)

    layer_rows = [
        ("L1", "輸入層", "基地 / 容積 / 獎勵拆解 / 免計基準 / 成本 / 都更全案投報參數（Sidebar）"),
        ("L2", "容積計算", "FA = 基地使用面積 × 容積率；允建 = FA × (1+獎勵率) + 容積移轉"),
        ("L3", "三項免計查核",
         "梯廳超出 = Σ max(0, 梯廳 − 樓板 × 梯廳%)，逐層；陽台同法；安全梯上限 = 允建 × 15%"),
        ("L4", "容積帳",
         "計入容積 = 圖說計入 + 梯廳超出 + 陽台超出；容積餘量 = 允建 − 計入（負值 = 超出）"),
        ("L4.5", "銷售坪效",
         "銷售坪 = (允建+陽台免計) × 外皮係數 / 3.3058 / (1−公設比)；銷坪比正常 1.58–1.68"),
        ("L5", "開發評效",
         "評效 = 總銷 / 總成本；> 5 優良 / 2–5 可行 / < 2 偏低（永盛內部定義）"),
        ("L6", "都更全案投報",
         "總銷 → 共同負擔六大科目（A工程/B管維/C權變/D利息/E稅捐/F管理）→ 共負比/分回/報酬率"),
        ("L7", "更新前估值 (P1)",
         "土地現值 + 建物殘值（RC 折舊 2%/年，§56）→ 增值倍率 = 地主分回 ÷ 更新前總值"),
    ]
    make_table(doc, ["層", "名稱", "核心邏輯"], layer_rows, col_widths=[1.0, 3.0, 12.5])

    make_section_heading(doc, "3.2  模組結構（v4.5 P0 模組化後）", level=2)

    module_rows = [
        ("app.py", "1,170 行", "UI 層（純 Streamlit），無計算邏輯；Sidebar / Tab / KPI / 圖表 / 報告匯出"),
        ("calc_engine.py", "446 行", "計算層（純函式，零 Streamlit 依賴）；12 個 calc_* 函式 + 範本資料 + 解析/報告"),
        ("law_db.py", "124 行", "法規庫；BONUS_都更(8項) / BONUS_危老(6項) / EXEMPTION_162 / COMMON_BURDEN_RANGES"),
        ("test_golden.py", "85 行", "黃金迴歸測試；4 案期望值鎖定，每次改公式必跑"),
    ]
    make_table(doc, ["模組", "規模", "職責"], module_rows, col_widths=[3.5, 2.0, 11.0])

    doc.add_paragraph()
    add_body(doc,
        "▶  模組化設計使計算函式可直接被 test_golden.py 引用，不需啟動 Streamlit；"
        "未來亦可作為獨立 Python package 供都更儀表板等下游工具 import。",
        color=MID_GRAY, size=9.5)

    make_section_heading(doc, "3.3  法規資料庫（law_db.py）", level=2)

    law_rows = [
        ("BONUS_都更", "8 項", "都市更新建築容積獎勵辦法；更新基本 ≤50% / 防災 ≤50% / 耐震 ≤10% / 綠建築 ≤10% / 智慧 ≤10% / 規模 ≤15% / 時程 ≤10% / TOD ≤20%"),
        ("BONUS_危老", "6 項", "危老條例 §6；基本 ≤10% / 規模 ≤10% / 時程 ≤30% / 耐震 ≤10% / 綠建築 ≤10% / 智慧 ≤5%"),
        ("EXEMPTION_162", "3 項", "建築技術規則 §162；梯廳逐層 5%/8%、安全梯 15%（允建基準）、陽台逐層 10%/15%"),
        ("COMMON_BURDEN_RANGES", "3 模式", "共負合理區間：都更全案管理 30–50% / 都更合建 35–55% / 危老重建 25–45%（超 65% 警示）"),
    ]
    make_table(doc, ["資料表", "項目數", "內容摘要"], law_rows, col_widths=[3.8, 1.5, 11.2])

    # ══════════════════════════════════════════════════════════════════
    # 4. 驗證與測試
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "4  驗證與品質保證", level=1)

    make_section_heading(doc, "4.1  黃金測試結果（test_golden.py — 最新執行 2026-06-23）", level=2)

    golden_rows = [
        ("安和段（都市更新）",
         "FA 4,569.71 m²\n允建 7,768.51 m²\n安全梯上限 1,165.28 m²",
         ("PASS ✅", True, GREEN_OK)),
        ("龜山半嶺段（危老）",
         "允建 4,244.04 m²\n梯廳超出 32.45 m²\n計入 4,276.25 m²\n餘量 −32.21 m²",
         ("PASS ✅", True, GREEN_OK)),
        ("中正段（防災都更）",
         "FA 2,211.75 m²\n允建 4,167.08 m²\n陽台超出 29.11 m²（逐層法）\n計入 4,198.75 m²\n餘量 −31.67 m²",
         ("PASS ✅", True, GREEN_OK)),
        ("中正段 L6 投報",
         "共負比 37.5%\n報酬率 166.3%\n共同負擔 108,803 萬",
         ("PASS ✅", True, GREEN_OK)),
    ]
    make_table(doc, ["測試案件", "核心期望值（容差 0.5 m²）", "結果"],
               golden_rows, col_widths=[4, 9.5, 3])

    make_section_heading(doc, "4.2  已修正重大問題（v3 起）", level=2)

    bug_rows = [
        ("陽台超出算法錯誤", "中正段陽台超出由 FA 總量法（303 m²）修正為逐層法（29.11 m²），"
         "差距約 10 倍。舊法高估超出量將誤判容積超出嚴重性。",
         ("已修正 v3", True, GREEN_OK)),
        ("中正段獎勵率偏差", "由原 0.6764 修正為 0.88407（防災都更 +88.4%），"
         "對應「獎勵 1,955.33 ÷ FA 2,211.75」精確反推。",
         ("已修正 v3", True, GREEN_OK)),
        ("面積表計入容積優先邏輯", "新增 面積表計入容積 參數：有值時直接採圖說彙總，避免逐層累加誤差。",
         ("已修正 v4.4", True, GREEN_OK)),
    ]
    make_table(doc, ["問題", "說明", "狀態"], bug_rows, col_widths=[3.5, 11, 2.0])

    # ══════════════════════════════════════════════════════════════════
    # 5. 版本演進
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "5  版本演進紀錄", level=1)

    ver_rows = [
        ("v4.1", "2026-06-10", "UI 精化；HTML KPI / Banner；都更全案投報 L6 上線"),
        ("v4.2", "2026-06-12", "設計感升級：藍圖 Hero、L2→L6 流程帶、KPI 進度條"),
        ("v4.3", "2026-06-17", "步驟化引導；§162 欄位對照卡；Windows 編碼修正"),
        ("v4.4", "2026-06-20", "面積表匯入優化；§162 核對表；CLAUDE.md 建築師反饋協議"),
        ("v4.5", "2026-06-22", "P0 模組化：calc_engine.py + law_db.py；獎勵拆解 UI（都更8/危老6項）；check_bonus_limit()"),
        ("v4.6", "2026-06-23", "P1 穩定：共負比合理區間警示；L7 更新前估值；增值倍率指標"),
    ]
    make_table(doc, ["版本", "日期", "主要變更"], ver_rows, col_widths=[1.5, 2.5, 12.5])

    # ══════════════════════════════════════════════════════════════════
    # 6. 已知問題與風險
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "6  已知問題與風險", level=1)

    issue_rows = [
        ("🔴 高", "安和段範本容積帳為空",
         "範本樓層表全為 0（安全梯/梯廳/陽台均空），導致容積帳計入=0、銷坪比偏低至 1.53（應為 1.62）。"
         "根因：缺少逐層面積種子資料。待提供安和段面積表後修正。"),
        ("🟠 中", "L7 增值倍率無防呆",
         "若地價輸入不合理，增值倍率可高達 7–11×（真實都更約 1.5–2.5×）。"
         "需加合理區間警示（建議：>3× 黃燈，>5× 紅燈）。"),
        ("🟠 中", "共負比對營造坪基準高度敏感",
         "中正段同一案，用「銷售坪」vs「允建坪」作為營造基準，共負比差 4 個百分點、報酬率差 33%。"
         "UI 有切換鈕但未在共負比卡片標示當前基準，易造成混淆。"),
        ("🟡 低", "版本號散落多處",
         "v4.6 字串分散在 app.py 封面橫幅、頁尾、calc_engine.py 報告函式，更版時易遺漏。"
         "建議抽成單一 VERSION 常數。"),
        ("🟡 低", "待建築師確認兩個法規參數",
         "梯廳免計基準 5% 還是 8%？陽台免計 10%（§162）還是 15%（§162-3）？"
         "數值不同將影響容積餘量 −2.49 到 −41.81 m²，需建築師確認條文依據後鎖定。"),
    ]
    make_table(doc, ["嚴重度", "問題", "說明"], issue_rows, col_widths=[1.5, 4, 11])

    # ══════════════════════════════════════════════════════════════════
    # 7. 後續建議執行展望
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "7  後續建議執行展望", level=1)

    make_section_heading(doc, "7.1  P1 完成：穩定現有資料（建議近期 2–3 週）", level=2)

    p1_rows = [
        ("P1-A", "補安和段逐層面積種子資料",
         "修正範本最大資料洞；將安和銷坪比 1.62 鎖進黃金測試第 5 案",
         "待提供安和段面積表", "🔴 高"),
        ("P1-B", "安和段共負費率校準",
         "用送審版六科目金額反推校準 財務率預設；新增安和共負比黃金測試",
         "待提供共負投報數據", "🔴 高"),
        ("P1-C", "L7 增值倍率合理區間",
         "比照共負比，加 >3× 黃燈 / >5× 紅燈防呆；附帶注釋提示地價輸入準確性",
         "無需外部資料，可立即實作", "🟠 中"),
        ("P1-D", "VERSION 常數化 + 報告版本更新",
         "抽出單一 VERSION = 'v4.6' 常數；同步更新 產生報告() 中的 v4.5 字串",
         "純程式清理，低風險", "🟡 低"),
    ]
    make_table(doc, ["項目", "工作", "說明", "前置條件", "優先度"],
               p1_rows, col_widths=[1.2, 3.5, 7.0, 3.5, 1.3])

    make_section_heading(doc, "7.2  P2 新功能：權利變換框架（建議 1–2 個月）", level=2)

    p2_rows = [
        ("P2-A", "地主清冊與權值模型",
         "各地主持分比例 × 土地/建物現值 → 更新前權利價值；為逐戶分回的分子"),
        ("P2-B", "逐戶分回模擬",
         "權值 × 地主分回總值 → 各戶應分回坪數；選配模組（大坪數 or 小坪數換多戶）"),
        ("P2-C", "找補金試算",
         "各戶分回應得值 vs 實際分配戶型市值 → 差額找補；符合都更條例 §51 找補機制"),
    ]
    make_table(doc, ["項目", "功能", "說明"], p2_rows, col_widths=[1.2, 3.5, 11.8])

    doc.add_paragraph()
    add_body(doc,
        "▶  P2 完成後，工具將覆蓋「可行性評估 → 概念設計 → 權利變換準備」全流程，"
        "可直接取代目前散落各 Excel 的 A 群（清冊/權值）+ C 群（投報主表）分頁。",
        color=MID_GRAY, size=9.5)

    make_section_heading(doc, "7.3  P3 整合：與都更儀表板（Urban-Renewal）結合", level=2)

    add_body(doc,
        "目標：RE-DCF-Tool（前期評估）的輸出作為 Urban-Renewal 儀表板（執行管理）的輸入，"
        "以「案件 JSON 合約」格式串接，讓一個案子從「評估通過」到「立案追蹤」零手動複製。",
        size=10)

    p3_rows = [
        ("A（近）", "匯出案件 JSON",
         "RE-DCF-Tool 新增「下載案件 JSON」按鈕，包含容積帳/坪效/投報/L7 估值，附 schema_version",
         "2 週", "無"),
        ("B（中）", "calc_engine 獨立套件化",
         "抽成 pip package（redcf-core），兩個 repo 共用計算核心；改公式一次同步兩端",
         "1 個月", "需確認儀表板技術棧"),
        ("C（長）", "雙向資料流",
         "儀表板真實權利人資料回傳 RE-DCF-Tool → 以實際清冊重算各地主分回比例",
         "3–6 個月", "需儀表板 API/DB 介面"),
    ]
    make_table(doc, ["階段", "工作", "說明", "估時", "前置"],
               p3_rows, col_widths=[1.2, 3.0, 8.5, 1.5, 2.3])

    make_section_heading(doc, "7.4  P4 長期：商辦版擴充", level=2)
    add_body(doc,
        "住宅版工具架構穩定後，複製計算層架構、替換為商辦財務模型（NOI + Cap Rate + 持有期報酬），"
        "另開對話與專案處理。預計複用 80% 的 UI 框架與測試架構，主要改寫 L5–L6 計算層。",
        size=10)

    # ══════════════════════════════════════════════════════════════════
    # 8. 待收資料清單
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "8  待收資料清單（安和段）", level=1)

    add_body(doc,
        "以下資料收到後可立即啟動 P1-A / P1-B 實作。標 ⭐ 為最高優先。",
        size=10, color=MID_GRAY)

    data_rows = [
        ("⭐ A1", "逐層面積表",
         "每層：樓板 / 計容積 / 梯廳 / 安全梯 / 陽台（m²）",
         "修正容積帳 + 坪效；鎖黃金測試"),
        ("⭐ B1", "共同負擔六科目金額",
         "A工程 / B管維 / C權變 / D利息 / E稅捐 / F管理（萬元）",
         "校準財務率預設；新增黃金測試"),
        ("⭐ B2", "總銷拆分",
         "住宅坪數&單價 / 店舖坪數&單價 / 車位數&單價",
         "校準 calc_總銷"),
        ("⭐ B3", "共負比 / 報酬率（送審定稿）",
         "最終核定數字，作為黃金測試期望值",
         "鎖定 L6 回歸"),
        ("B4", "營造坪基準",
         "是銷售坪還是允建坪（影響 4 個百分點共負比）",
         "確保與送審口徑一致"),
        ("C1", "更新前土地單價或總現值",
         "萬/坪 或 萬元（估價師核定或地價標準）",
         "校準 L7 增值倍率"),
        ("A2*", "梯廳 / 陽台免計基準確認",
         "建築師確認：梯廳 5% 或 8%？陽台 10% 或 §162-3 的 15%？",
         "法規基準確定後鎖定"),
    ]
    make_table(doc, ["優先", "資料", "明細", "用途"], data_rows,
               col_widths=[1.2, 3.0, 7.5, 4.8])

    # ══════════════════════════════════════════════════════════════════
    # 附錄：法規背景（含 research_notes）
    # ══════════════════════════════════════════════════════════════════
    make_section_heading(doc, "附錄 A  法規背景查證摘要", level=1)

    if research_notes and research_notes.strip():
        add_body(doc, research_notes, size=9.5, color=DARK_NAVY)
    else:
        add_body(doc,
            "以下為已知法規要點（查證進行中，結果將補充）：",
            size=10, color=MID_GRAY)
        law_notes = [
            "都市更新條例（2019 年大幅修正，2023 年小修）：§56 規定以更新前各宗土地及建物現值"
            "為權利變換計算基礎，RC 造耐用年限 50 年、折舊率 2%/年為常見標準。",
            "都市更新建築容積獎勵辦法（依都更條例 §65 授權）：更新基本獎勵上限 FA×50%（§3）；"
            "防災都更額外最高 50%（§65 第4項）；耐震/綠建築/智慧/規模/時程/TOD 各有上限。",
            "危老重建條例（2017年施行，已數次延長申請期限）：§6 時程獎勵最高 30%（第1年）；"
            "各縣市另訂容積放寬上限（台北市合計上限依個案而定）。",
            "建築技術規則 §162：三項免計面積（梯廳/機電/陽台）均為「逐層」計算，"
            "非以 FA 或允建容積總量乘以比率——此點為本工具核心設計依據。",
            "Streamlit Community Cloud 免費方案（2025–2026 年現行）：無限 public app 部署；"
            "單 app RAM 限制約 1 GB；付費方案（Teams/Enterprise）另計。",
        ]
        for n in law_notes:
            add_bullet(doc, n, color=DARK_NAVY)
        add_body(doc,
            "\n* 以上法規資訊以主管機關公告及全國法規資料庫為準，如有修正請以最新版本為據。",
            size=9, color=LIGHT_GRAY, italic=True)

    # 頁尾
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(footer_p, color="E7E9F2", sz=4)
    r = footer_p.add_run(
        "RE-DCF-Tool　｜　永盛開發建設前期評估工具　｜　v4.6　｜　build 2026-06-23"
        "　｜　本報告由 generate_report.py 自動生成")
    r.font.size = Pt(8)
    r.font.color.rgb = LIGHT_GRAY

    # 儲存
    out_path = "/home/user/RE-DCF-Tool/RE-DCF-Tool_專案報告_2026-06-26.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build_report()
    print(f"✅ 報告已生成：{path}")
